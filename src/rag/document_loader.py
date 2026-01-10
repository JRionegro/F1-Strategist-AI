"""
Document Loader for RAG System.

Loads and processes documents from various formats (MD, PDF, DOCX) for
indexing into ChromaDB. Supports year/circuit-based document organization.
"""

import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import yaml

logger = logging.getLogger(__name__)


@dataclass
class Document:
    """Represents a document chunk for RAG indexing."""

    content: str
    metadata: dict = field(default_factory=dict)
    doc_id: str = ""

    def __post_init__(self):
        """Generate doc_id if not provided."""
        if not self.doc_id:
            # Generate unique ID from source, chunk index and full content hash
            source = self.metadata.get("source", "unknown")
            chunk_index = self.metadata.get("chunk_index", 0)
            # Use full content hash for uniqueness
            import hashlib
            content_hash = hashlib.md5(self.content.encode()).hexdigest()[:8]
            self.doc_id = f"{source}_chunk{chunk_index}_{content_hash}"


@dataclass
class DocumentInfo:
    """Information about an available document."""

    filename: str
    path: str
    category: str
    format: str  # md, pdf, docx
    year: Optional[int] = None
    circuit: Optional[str] = None
    tags: list[str] = field(default_factory=list)


class DocumentLoader:
    """
    Load and chunk documents for RAG indexing.

    Supports:
    - Markdown (.md) - Direct processing
    - PDF (.pdf) - Converted to text/markdown
    - Word (.docx) - Converted to text/markdown

    Directory structure:
    data/rag/
    ├── global/           # Always loaded
    ├── templates/        # Not loaded, used for generation
    ├── {year}/          # Year-specific docs
    │   ├── *.md         # Year-level docs (regulations, compounds)
    │   └── circuits/
    │       └── {circuit}/ # Circuit-specific docs
    """

    # Supported file extensions
    SUPPORTED_FORMATS = {".md", ".pdf", ".docx", ".doc"}

    # Default chunk settings
    DEFAULT_CHUNK_SIZE = 1000
    DEFAULT_CHUNK_OVERLAP = 200

    def __init__(
        self,
        base_path: str = "data/rag",
        chunk_size: int = DEFAULT_CHUNK_SIZE,
        chunk_overlap: int = DEFAULT_CHUNK_OVERLAP,
    ):
        """
        Initialize DocumentLoader.

        Args:
            base_path: Base directory for RAG documents
            chunk_size: Target size for document chunks (characters)
            chunk_overlap: Overlap between chunks (characters)
        """
        self.base_path = Path(base_path)
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

        # Ensure base path exists
        self.base_path.mkdir(parents=True, exist_ok=True)

        logger.info(
            f"DocumentLoader initialized with base_path={base_path}, "
            f"chunk_size={chunk_size}, overlap={chunk_overlap}"
        )

    def load_documents_for_context(
        self,
        year: int,
        circuit: Optional[str] = None,
    ) -> list[Document]:
        """
        Load documents for a specific year and optionally circuit.

        Loads in order:
        1. global/*.md (always)
        2. {year}/*.md (year-level docs)
        3. {year}/circuits/{circuit}/*.md (if circuit specified)

        Args:
            year: Target year (e.g., 2024)
            circuit: Circuit name in lowercase (e.g., "bahrain", "monaco")

        Returns:
            List of Document objects ready for indexing
        """
        documents: list[Document] = []

        # 1. Load global documents
        global_path = self.base_path / "global"
        if global_path.exists():
            docs = self._load_directory(global_path, {"scope": "global"})
            documents.extend(docs)
            logger.info(f"Loaded {len(docs)} global documents")

        # 2. Load year-level documents
        year_path = self.base_path / str(year)
        if year_path.exists():
            # Load only files in year root (not circuits subfolder)
            year_docs = self._load_directory(
                year_path,
                {"year": year, "scope": "year"},
                recursive=False,
            )
            documents.extend(year_docs)
            logger.info(f"Loaded {len(year_docs)} documents for year {year}")

        # 3. Load circuit-specific documents
        if circuit:
            circuit_path = year_path / "circuits" / circuit.lower()
            if circuit_path.exists():
                circuit_docs = self._load_directory(
                    circuit_path,
                    {"year": year, "circuit": circuit.lower(), "scope": "circuit"},
                )
                documents.extend(circuit_docs)
                logger.info(
                    f"Loaded {len(circuit_docs)} documents for {circuit} {year}"
                )
            else:
                logger.warning(
                    f"No documents found for circuit {circuit} in {year}"
                )

        logger.info(
            f"Total documents loaded: {len(documents)} "
            f"(year={year}, circuit={circuit})"
        )
        return documents

    def _load_directory(
        self,
        directory: Path,
        base_metadata: dict,
        recursive: bool = True,
    ) -> list[Document]:
        """
        Load all supported documents from a directory.

        Args:
            directory: Directory path to load from
            base_metadata: Metadata to add to all documents
            recursive: Whether to recurse into subdirectories

        Returns:
            List of Document objects
        """
        documents: list[Document] = []

        if not directory.exists():
            return documents

        # Get files based on recursive setting
        if recursive:
            files = list(directory.rglob("*"))
        else:
            files = list(directory.glob("*"))

        for file_path in files:
            if not file_path.is_file():
                continue

            suffix = file_path.suffix.lower()
            if suffix not in self.SUPPORTED_FORMATS:
                continue

            try:
                # Load and process file
                docs = self._load_file(file_path, base_metadata)
                documents.extend(docs)
            except Exception as e:
                logger.error(f"Error loading {file_path}: {e}")

        return documents

    def _load_file(
        self,
        file_path: Path,
        base_metadata: dict,
    ) -> list[Document]:
        """
        Load a single file and convert to Document chunks.

        Args:
            file_path: Path to the file
            base_metadata: Base metadata to include

        Returns:
            List of Document chunks
        """
        suffix = file_path.suffix.lower()

        # Extract content based on format
        if suffix == ".md":
            content, frontmatter = self._load_markdown(file_path)
        elif suffix == ".pdf":
            content = self._load_pdf(file_path)
            frontmatter = {}
        elif suffix in {".docx", ".doc"}:
            content = self._load_docx(file_path)
            frontmatter = {}
        else:
            logger.warning(f"Unsupported format: {suffix}")
            return []

        if not content or not content.strip():
            logger.warning(f"Empty content from {file_path}")
            return []

        # Determine category for this file
        category = self._categorize_file(file_path)

        # Build metadata
        metadata = {
            **base_metadata,
            "source": str(file_path.relative_to(self.base_path)),
            "filename": file_path.name,
            "format": suffix[1:],  # Remove dot
            "category": category,
            **frontmatter,
        }

        # Chunk the content
        chunks = self._chunk_content(content, metadata)

        logger.debug(
            f"Loaded {file_path.name}: {len(content)} chars -> {len(chunks)} chunks"
        )

        return chunks

    def _load_markdown(self, file_path: Path) -> tuple[str, dict]:
        """
        Load markdown file, extracting YAML frontmatter if present.

        Args:
            file_path: Path to markdown file

        Returns:
            Tuple of (content, frontmatter_dict)
        """
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        frontmatter = {}

        # Check for YAML frontmatter (--- delimited)
        if content.startswith("---"):
            parts = content.split("---", 2)
            if len(parts) >= 3:
                try:
                    frontmatter = yaml.safe_load(parts[1]) or {}
                    content = parts[2].strip()
                except yaml.YAMLError as e:
                    logger.warning(f"Invalid YAML frontmatter in {file_path}: {e}")

        return content, frontmatter

    def _load_pdf(self, file_path: Path) -> str:
        """
        Load PDF file and convert to text.

        Uses pymupdf4llm if available, falls back to basic extraction.

        Args:
            file_path: Path to PDF file

        Returns:
            Extracted text content
        """
        try:
            # Try pymupdf4llm for better markdown conversion
            import pymupdf4llm  # type: ignore[import-untyped]

            md_text = pymupdf4llm.to_markdown(str(file_path))
            logger.info(f"Loaded PDF with pymupdf4llm: {file_path.name}")
            return str(md_text)

        except ImportError:
            logger.warning(
                "pymupdf4llm not installed. Install with: pip install pymupdf4llm"
            )

            # Fallback to basic PyMuPDF extraction
            try:
                import fitz  # type: ignore[import-untyped]  # PyMuPDF

                doc = fitz.open(file_path)
                text_parts = []
                for page in doc:
                    text_parts.append(page.get_text())
                doc.close()

                logger.info(f"Loaded PDF with PyMuPDF: {file_path.name}")
                return "\n\n".join(text_parts)

            except ImportError:
                logger.error(
                    "PyMuPDF not installed. Install with: pip install pymupdf"
                )
                return ""

    def _load_docx(self, file_path: Path) -> str:
        """
        Load Word document and convert to text.

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text content
        """
        try:
            from docx import Document as DocxDocument  # type: ignore[import-untyped]

            doc = DocxDocument(str(file_path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            logger.info(f"Loaded DOCX: {file_path.name}")
            return "\n\n".join(paragraphs)

        except ImportError:
            logger.error(
                "python-docx not installed. Install with: pip install python-docx"
            )
            return ""

    def _chunk_content(
        self,
        content: str,
        metadata: dict,
    ) -> list[Document]:
        """
        Split content into chunks suitable for embedding.

        Uses markdown-aware splitting when possible.

        Args:
            content: Full document content
            metadata: Metadata to attach to each chunk

        Returns:
            List of Document chunks
        """
        chunks: list[Document] = []

        # Try to use LangChain splitter if available
        try:
            from langchain_text_splitters import RecursiveCharacterTextSplitter  # type: ignore[import-untyped]

            # Markdown-aware separators
            separators = [
                "\n## ",      # H2 headers
                "\n### ",     # H3 headers
                "\n#### ",    # H4 headers
                "\n\n",       # Paragraphs
                "\n",         # Lines
                ". ",         # Sentences
                " ",          # Words
                "",           # Characters
            ]

            splitter = RecursiveCharacterTextSplitter(
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap,
                separators=separators,
                length_function=len,
            )

            texts = splitter.split_text(content)

            for i, text in enumerate(texts):
                chunk_metadata = {
                    **metadata,
                    "chunk_index": i,
                    "total_chunks": len(texts),
                }
                chunks.append(Document(content=text, metadata=chunk_metadata))

        except ImportError:
            # Fallback to simple splitting
            logger.warning(
                "LangChain not installed, using simple chunking. "
                "Install with: pip install langchain-text-splitters"
            )
            chunks = self._simple_chunk(content, metadata)

        return chunks

    def _simple_chunk(
        self,
        content: str,
        metadata: dict,
    ) -> list[Document]:
        """
        Simple fallback chunking by paragraphs.

        Args:
            content: Full content
            metadata: Base metadata

        Returns:
            List of Document chunks
        """
        chunks: list[Document] = []

        # Split by double newlines (paragraphs)
        paragraphs = re.split(r"\n\n+", content)

        current_chunk = ""
        chunk_index = 0

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            # Check if adding this paragraph exceeds chunk size
            if len(current_chunk) + len(para) > self.chunk_size:
                if current_chunk:
                    chunk_metadata = {**metadata, "chunk_index": chunk_index}
                    chunks.append(
                        Document(content=current_chunk.strip(), metadata=chunk_metadata)
                    )
                    chunk_index += 1

                # Start new chunk with overlap from end of previous
                overlap_text = current_chunk[-self.chunk_overlap:] if current_chunk else ""
                current_chunk = overlap_text + "\n\n" + para if overlap_text else para
            else:
                current_chunk = (
                    current_chunk + "\n\n" + para if current_chunk else para
                )

        # Add final chunk
        if current_chunk.strip():
            chunk_metadata = {**metadata, "chunk_index": chunk_index}
            chunks.append(
                Document(content=current_chunk.strip(), metadata=chunk_metadata)
            )

        return chunks

    def get_available_documents(
        self,
        year: Optional[int] = None,
        circuit: Optional[str] = None,
    ) -> dict[str, list[DocumentInfo]]:
        """
        List available documents organized by category.

        Args:
            year: Filter by year (optional)
            circuit: Filter by circuit (optional)

        Returns:
            Dict mapping category to list of DocumentInfo
        """
        documents: dict[str, list[DocumentInfo]] = {
            "global": [],
            "fia": [],
            "strategy": [],
            "weather": [],
            "performance": [],
            "race_control": [],
            "race_position": [],
            "other": [],
        }

        # Scan global
        global_path = self.base_path / "global"
        if global_path.exists():
            self._scan_directory(global_path, documents, scope="global")

        # Scan year if specified
        if year:
            year_path = self.base_path / str(year)
            if year_path.exists():
                self._scan_directory(
                    year_path, documents, year=year, recursive=False
                )

                # Scan circuit if specified
                if circuit:
                    circuit_path = year_path / "circuits" / circuit.lower()
                    if circuit_path.exists():
                        self._scan_directory(
                            circuit_path, documents, year=year, circuit=circuit
                        )

        return documents

    def _scan_directory(
        self,
        directory: Path,
        documents: dict[str, list[DocumentInfo]],
        year: Optional[int] = None,
        circuit: Optional[str] = None,
        scope: str = "year",
        recursive: bool = True,
    ) -> None:
        """
        Scan directory and categorize documents.

        Args:
            directory: Directory to scan
            documents: Dict to populate with DocumentInfo
            year: Year context
            circuit: Circuit context
            scope: Scope label
            recursive: Whether to recurse
        """
        if recursive:
            files = list(directory.rglob("*"))
        else:
            files = list(directory.glob("*"))

        for file_path in files:
            if not file_path.is_file():
                continue

            suffix = file_path.suffix.lower()
            if suffix not in self.SUPPORTED_FORMATS:
                continue

            # Determine category from filename or frontmatter
            category = self._categorize_file(file_path)

            doc_info = DocumentInfo(
                filename=file_path.name,
                path=str(file_path.relative_to(self.base_path)),
                category=category,
                format=suffix[1:],
                year=year,
                circuit=circuit,
            )

            if category in documents:
                documents[category].append(doc_info)
            else:
                documents["other"].append(doc_info)

    def _categorize_file(self, file_path: Path) -> str:
        """
        Determine document category from filename or content.

        Args:
            file_path: Path to file

        Returns:
            Category string
        """
        filename = file_path.name.lower()

        # Check filename patterns
        if "regulation" in filename or "fia" in filename:
            return "fia"
        if "strategy" in filename or "pit" in filename:
            return "strategy"
        if "weather" in filename or "temperature" in filename:
            return "weather"
        if "performance" in filename:
            return "performance"
        if "race_control" in filename or "racecontrol" in filename:
            return "race_control"
        if "race_position" in filename or "position" in filename:
            return "race_position"
        if file_path.parent.name == "global":
            return "global"

        # Try to read frontmatter for category
        if file_path.suffix.lower() == ".md":
            try:
                _, frontmatter = self._load_markdown(file_path)
                if "category" in frontmatter:
                    return frontmatter["category"]
            except Exception:
                pass

        return "other"

    def convert_pdf_to_markdown(
        self,
        pdf_path: str,
        output_path: Optional[str] = None,
    ) -> str:
        """
        Convert a PDF file to Markdown and optionally save it.

        Useful for converting FIA regulations PDFs to indexable format.

        Args:
            pdf_path: Path to source PDF
            output_path: Optional path to save MD file

        Returns:
            Markdown content
        """
        pdf_file = Path(pdf_path)
        if not pdf_file.exists():
            raise FileNotFoundError(f"PDF not found: {pdf_path}")

        content = self._load_pdf(pdf_file)

        if output_path:
            output_file = Path(output_path)
            output_file.parent.mkdir(parents=True, exist_ok=True)

            # Add metadata header
            metadata_header = f"""---
category: fia
source_pdf: {pdf_file.name}
converted: true
---

"""
            with open(output_file, "w", encoding="utf-8") as f:
                f.write(metadata_header + content)

            logger.info(f"Saved converted PDF to {output_path}")

        return content

    def convert_docx_to_markdown(
        self,
        docx_path: str,
        output_path: Optional[str] = None,
    ) -> str:
        """
        Convert a DOCX file to Markdown and optionally save it.

        Args:
            docx_path: Path to source DOCX
            output_path: Optional path to save MD file

        Returns:
            Markdown content
        """
        docx_file = Path(docx_path)
        if not docx_file.exists():
            raise FileNotFoundError(f"DOCX not found: {docx_path}")

        content = self._load_docx(docx_file)

        if output_path:
            output_file = Path(output_path)
            output_file.parent.mkdir(parents=True, exist_ok=True)

            # Add metadata header
            metadata_header = f"""---
source_docx: {docx_file.name}
converted: true
---

"""
            with open(output_file, "w", encoding="utf-8") as f:
                f.write(metadata_header + content)

            logger.info(f"Saved converted DOCX to {output_path}")

        return content

    def suggest_category_with_llm(
        self,
        document_content: str,
        filename: str,
        llm_manager=None
    ) -> dict:
        """
        Suggest document category using LLM analysis.
        
        Args:
            document_content: Full document text or excerpt
            filename: Original filename for context
            llm_manager: LLM manager instance (optional, will import globally)
            
        Returns:
            Dict with keys: category, confidence, reasoning
        """
        try:
            # Extract first 2000 characters for analysis
            excerpt = document_content[:2000] if len(document_content) > 2000 else document_content
            
            # Import LLM manager if not provided
            if llm_manager is None:
                try:
                    from src.llm.llm_manager import LLMManager
                    llm_manager = LLMManager()
                except Exception as e:
                    logger.warning(f"Could not import LLM manager: {e}")
                    return {
                        'category': 'other',
                        'confidence': 0.0,
                        'reasoning': 'LLM unavailable - manual categorization required'
                    }
            
            # Build prompt
            prompt = f"""Analyze this F1 document and categorize it into ONE of these categories:

Categories:
- strategy: Race strategy, pit stop analysis, tire management, race planning
- weather: Weather conditions, forecasts, track temperature, meteorological data
- performance: Lap times, telemetry data, car performance, sector times
- race_control: Flags, penalties, safety car, VSC, race incidents, steward decisions
- race_position: Position tracking, overtakes, gaps between drivers, race order
- fia: FIA regulations, sporting code, technical regulations, rule clarifications
  * Examples: "FIA Formula 1 Sporting Regulations", "FIA Technical Regulations", 
    "International Sporting Code", documents with "Article", "regulation", "sporting code"
  * Keywords: "FIA", "sporting regulations", "technical regulations", "article", 
    "appendix", "championship", "regulatory", "official", "federation"
- global: General F1 information, circuit layouts, driver info, team data

IMPORTANT: If the document contains official FIA regulations, sporting code, or technical rules 
with numbered articles and formal language, it should ALWAYS be categorized as 'fia', even if 
it also discusses strategy or performance.

Document Information:
- Filename: {filename}
- Content excerpt (first 2000 chars):
{excerpt}

Analyze the content and return a JSON object with:
{{
    "category": "one of the categories above",
    "confidence": 0.0 to 1.0 (how certain you are),
    "reasoning": "brief explanation of why this category fits"
}}

Only return the JSON object, no other text."""

            # Call LLM (try Gemini first, fallback to OpenAI)
            try:
                response = llm_manager.generate(
                    prompt=prompt,
                    provider="gemini",
                    temperature=0.3,
                    max_tokens=200
                )
            except Exception:
                try:
                    response = llm_manager.generate(
                        prompt=prompt,
                        provider="openai",
                        temperature=0.3,
                        max_tokens=200
                    )
                except Exception as e2:
                    logger.warning(f"LLM call failed: {e2}")
                    return {
                        'category': 'other',
                        'confidence': 0.0,
                        'reasoning': 'LLM call failed - manual categorization required'
                    }
            
            # Parse JSON response
            import json
            try:
                # Remove markdown code blocks if present
                clean_response = response.strip()
                if clean_response.startswith('```'):
                    clean_response = clean_response.split('```')[1]
                    if clean_response.startswith('json'):
                        clean_response = clean_response[4:]
                clean_response = clean_response.strip()
                
                result = json.loads(clean_response)
                
                # Validate category
                valid_categories = ['strategy', 'weather', 'performance', 'race_control', 'race_position', 'fia', 'global', 'other']
                if result.get('category') not in valid_categories:
                    result['category'] = 'other'
                    result['confidence'] = 0.0
                
                # Ensure confidence is float
                result['confidence'] = float(result.get('confidence', 0.0))
                
                logger.info(f"LLM suggested category: {result['category']} (confidence: {result['confidence']:.2f})")
                return result
                
            except Exception as parse_error:
                logger.warning(f"Could not parse LLM response: {parse_error}")
                return {
                    'category': 'other',
                    'confidence': 0.0,
                    'reasoning': f'Could not parse LLM response: {str(parse_error)[:100]}'
                }
                
        except Exception as e:
            logger.error(f"Error in suggest_category_with_llm: {e}", exc_info=True)
            return {
                'category': 'other',
                'confidence': 0.0,
                'reasoning': f'Error during analysis: {str(e)[:100]}'
            }

    def backup_existing_document(self, filepath: Path) -> Optional[Path]:
        """
        Create backup of existing document before replacement.
        
        Args:
            filepath: Path to file that will be replaced
            
        Returns:
            Path to backup file, or None if file doesn't exist
        """
        try:
            filepath = Path(filepath)
            
            if not filepath.exists():
                logger.info(f"No existing file to backup: {filepath}")
                return None
            
            # Create backup filename with timestamp
            from datetime import datetime
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            backup_filename = f"{filepath.stem}.backup.{timestamp}{filepath.suffix}"
            backup_path = filepath.parent / backup_filename
            
            # Copy file to backup
            import shutil
            shutil.copy2(filepath, backup_path)
            
            logger.info(f"Created backup: {backup_path}")
            
            # Log to upload history
            self._log_backup_to_history(filepath, backup_path)
            
            return backup_path
            
        except Exception as e:
            logger.error(f"Error creating backup: {e}", exc_info=True)
            return None

    def _log_backup_to_history(self, original_path: Path, backup_path: Path):
        """
        Log backup creation to .upload_history.json.
        
        Args:
            original_path: Original file path
            backup_path: Backup file path
        """
        try:
            from datetime import datetime
            import json
            
            history_file = self.base_dir / ".upload_history.json"
            
            # Load existing history
            if history_file.exists():
                with open(history_file, 'r', encoding='utf-8') as f:
                    history = json.load(f)
            else:
                history = []
            
            # Add new entry
            entry = {
                'timestamp': datetime.now().isoformat(),
                'action': 'backup',
                'original_file': str(original_path.relative_to(self.base_dir)),
                'backup_file': str(backup_path.relative_to(self.base_dir))
            }
            history.append(entry)
            
            # Save history
            with open(history_file, 'w', encoding='utf-8') as f:
                json.dump(history, f, indent=2)
                
            logger.debug(f"Logged backup to history: {entry}")
            
        except Exception as e:
            logger.warning(f"Could not log backup to history: {e}")

